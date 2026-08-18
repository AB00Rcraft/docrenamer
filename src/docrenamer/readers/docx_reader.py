"""DOCX (раздел 17 ТЗ).

Извлекаются абзацы, таблицы, колонтитулы и core properties. Обратно ничего не
сохраняется.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

from docrenamer.readers.base import finalize_text, safe_metadata
from docrenamer.types import ReadResult, Status, nfc

if TYPE_CHECKING:  # pragma: no cover
    from docrenamer.analysis import ReaderContext

#: Ограничения обхода: документы бывают очень большими.
MAX_PARAGRAPHS = 4000
MAX_TABLE_CELLS = 4000


def read_docx(path: Path, context: ReaderContext) -> ReadResult:
    """Прочитать DOCX."""
    result = ReadResult()
    limits = context.limits
    try:
        import docx
    except ImportError:  # pragma: no cover
        result.add_status(Status.UNSUPPORTED_FORMAT)
        return result

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # недоверенный вход: python-docx поднимает разное
        result.add_status(Status.READ_ERROR)
        result.decoding_warnings.append(f"DOCX не открыт: {exc}")
        return result

    parts: list[str] = []
    for index, paragraph in enumerate(document.paragraphs):
        if index >= MAX_PARAGRAPHS:
            result.add_status(Status.LIMIT_EXCEEDED)
            break
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    cells = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cells += 1
                if cells > MAX_TABLE_CELLS:
                    result.add_status(Status.LIMIT_EXCEEDED)
                    break
                text = cell.text.strip()
                if text:
                    parts.append(text)
            if cells > MAX_TABLE_CELLS:
                break
        if cells > MAX_TABLE_CELLS:
            break

    headers_footers: list[str] = []
    try:
        for section in document.sections:
            for container in (section.header, section.footer):
                for paragraph in container.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        headers_footers.append(text)
    except (AttributeError, ValueError, KeyError):
        pass

    result.metadata.update(_core_properties(document))
    result.metadata.update(
        safe_metadata(
            {
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "headers_footers": headers_footers[:20],
            }
        )
    )
    result.source_encoding = "ooxml/utf-8"
    result.encoding_confidence = 1.0
    body = "\n".join(parts)
    if headers_footers:
        body = "\n".join([*headers_footers[:5], body])
    return finalize_text(result, body, limits)


def _core_properties(document: Any) -> dict[str, Any]:
    """Стандартные свойства OOXML-документа."""
    values: dict[str, Any] = {}
    try:
        core = document.core_properties
    except (AttributeError, KeyError):
        return values
    mapping = {
        "title": "title",
        "author": "author",
        "subject": "subject",
        "created": "created",
        "modified": "modified",
        "last_modified_by": "last_modified_by",
        "category": "category",
        "comments": "comments",
        "keywords": "keywords",
    }
    for attribute, name in mapping.items():
        try:
            value = getattr(core, attribute, None)
        except (AttributeError, ValueError):
            continue
        if value:
            values[name] = nfc(str(value))
    return safe_metadata(values)
